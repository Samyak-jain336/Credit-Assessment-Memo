"""
docx_writer.py

Creates the final Credit Assessment Memo (CAM) document.
Renders LLM-generated markdown content into properly formatted
Word paragraphs — bold, bullets, line breaks, and justified text.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from stubs import CAMSection


def _render_inline(paragraph, text: str):
    """Render a line of text with inline bold/italic markdown into
    Word runs. Handles **bold** markers. Strips backtick code spans
    (replaces with plain text). Skips literal asterisk markers that
    the LLM sometimes outputs as \\* or \*.
    """
    import re

    # Strip backtick code spans — render content as plain text
    text = re.sub(r'`([^`]+)`', r'\1', text)

    # Strip leading markdown heading markers (## ### etc.)
    text = re.sub(r'^#+\s*', '', text)

    # Skip if this is a duplicate section title line
    # (LLM often outputs "**Section: X**" as first line)
    if re.match(r'^\*\*Section:', text.strip()):
        return

    # Split on **bold** markers
    parts = re.split(r'\*\*', text)
    is_bold = False
    for part in parts:
        if part:
            run = paragraph.add_run(part)
            run.bold = is_bold
            run.font.size = Pt(11)
        is_bold = not is_bold


class DocumentWriter:
    """Responsible for writing the final CAM document."""

    def __init__(self):
        self.document = Document()
        self._set_default_style()

    def _set_default_style(self):
        """Set document-wide default font."""
        style = self.document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def _add_title(self, company_name: str, fiscal_year: int):
        heading = self.document.add_heading('Credit Assessment Memo', level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p = self.document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r1 = p.add_run('Company: ')
        r1.bold = True
        r1.font.size = Pt(12)
        p.add_run(company_name).font.size = Pt(12)

        p2 = self.document.add_paragraph()
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r2 = p2.add_run('Financial Year: ')
        r2.bold = True
        r2.font.size = Pt(12)
        p2.add_run(str(fiscal_year)).font.size = Pt(12)

        self.document.add_page_break()

    def _add_section(self, section: CAMSection):
        import re

        # Section heading
        heading = self.document.add_heading(section.title, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

        content = section.content

        # Normalise line endings
        content = content.replace('\r\n', '\n').replace('\r', '\n')

        # Split into lines
        lines = content.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i].strip()
            i += 1

            # Skip empty lines — Word already has paragraph spacing
            if not line:
                continue

            # Skip duplicate section title lines from LLM
            if re.match(r'^\*\*Section:', line):
                continue

            # Skip standalone markdown heading that repeats section title
            if re.match(r'^#+\s*' + re.escape(section.title), line, re.IGNORECASE):
                continue

            # Bullet point
            if line.startswith('* ') or line.startswith('- '):
                bullet_text = line[2:].strip()
                p = self.document.add_paragraph(style='List Bullet')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                _render_inline(p, bullet_text)
                continue

            # Subheading (## or ###)
            if line.startswith('#'):
                clean = re.sub(r'^#+\s*', '', line)
                p = self.document.add_paragraph()
                r = p.add_run(clean)
                r.bold = True
                r.font.size = Pt(11)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                continue

            # Normal paragraph
            p = self.document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            _render_inline(p, line)

    def _add_reconciliation_table(self, reconciliation_results):
        """Render reconciliation results as a Word table.
        Called only for the Data Consistency Review section.
        No LLM involved — renders directly from ReconciliationResult
        objects, which are structured Python data from reconciliation.py.
        """
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        hdr = table.rows[0].cells
        hdr[0].text = 'Field'
        hdr[1].text = 'Database Value (Cr)'
        hdr[2].text = 'Evidence Found'
        hdr[3].text = 'Status'
        hdr[4].text = 'Remarks'

        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        for result in reconciliation_results:
            row = table.add_row().cells
            row[0].text = result.field_name.replace('_', ' ').title()
            row[1].text = (
                str(result.database_value)
                if result.database_value is not None
                else 'N/A'
            )
            row[2].text = (
                str(result.extracted_value)
                if result.extracted_value
                else 'Not found'
            )
            row[3].text = result.status.upper()
            row[4].text = result.remarks if result.remarks else "—"

        self.document.add_paragraph()

    def _add_generic_table(self, table_data: list, columns: list):
        """Render LLM-generated structured table data as a Word table.
        Called after the narrative for Financial Analysis, Banking
        Conduct, and Risk Assessment when table_data is present on
        the CAMSection object.
        """
        if not table_data:
            return

        table = self.document.add_table(rows=1, cols=len(columns))
        table.style = 'Light Grid Accent 1'

        hdr = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr[i].text = col.replace('_', ' ').title()
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True

        for row_data in table_data:
            row = table.add_row().cells
            for i, col in enumerate(columns):
                value = row_data.get(col)
                row[i].text = str(value) if value is not None else '—'

        self.document.add_paragraph()

    def write(
        self,
        company_name: str,
        fiscal_year: int,
        sections: list,
        output_path: str,
        reconciliation_results=None,
    ) -> str:
        """
        Generate the final CAM document.

        Data Consistency Review is rendered as a code-generated table
        directly from reconciliation_results — no LLM prose used for
        that section. All other sections render narrative first, then
        an LLM-generated table if table_data is present on the section.
        """
        import schema

        self._add_title(company_name, fiscal_year)

        for section in sections:
            if section.title == "Data Consistency Review" and reconciliation_results:
                heading = self.document.add_heading(section.title, level=1)
                for run in heading.runs:
                    run.font.size = Pt(14)
                    run.font.bold = True
                self._add_reconciliation_table(reconciliation_results)
            else:
                self._add_section(section)
                if section.table_data:
                    columns = schema.SECTION_TABLE_SCHEMA.get(section.title)
                    if columns:
                        self._add_generic_table(section.table_data, columns)

        self.document.save(output_path)
        return output_path