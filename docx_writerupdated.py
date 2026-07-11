"""
docx_writer.py

Creates the final Credit Assessment Memo (CAM) document.

Responsibilities
----------------
- Create a Word document.
- Add title and company metadata.
- Insert generated CAM sections.
- Render reconciliation results as a table.
- Save the document.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from stubs import CAMSection, ReconciliationResult


class DocumentWriter:
    """
    Responsible for writing the final CAM document.
    """

    def __init__(self):
        self.document = Document()

    # ---------------------------------------------------------
    # Helpers
    # ---------------------------------------------------------

    def _add_title(
        self,
        company_name: str,
        fiscal_year: int,
    ):

        heading = self.document.add_heading(
            "Credit Assessment Memo",
            level=0,
        )

        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p = self.document.add_paragraph()

        p.add_run("Company: ").bold = True
        p.add_run(company_name)

        p.add_run("\nFinancial Year: ").bold = True
        p.add_run(str(fiscal_year))

        self.document.add_page_break()

    # ---------------------------------------------------------
    # Section Writer
    # ---------------------------------------------------------

    def _add_section(
        self,
        section: CAMSection,
    ):

        heading = self.document.add_heading(
            section.title,
            level=1,
        )

        heading.style.font.size = Pt(16)

        body = self.document.add_paragraph()

        body.style.font.size = Pt(11)

        body.add_run(section.content)

    # ---------------------------------------------------------
    # Reconciliation Table
    # ---------------------------------------------------------

    def _add_reconciliation_table(
        self,
        reconciliation_results: list[ReconciliationResult],
    ):
        """
        Render reconciliation results as a Word table.
        """

        if not reconciliation_results:
            self.document.add_paragraph(
                "No reconciliation issues found."
            )
            return

        table = self.document.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        headers = table.rows[0].cells
        headers[0].text = "Field"
        headers[1].text = "Database Value"
        headers[2].text = "Extracted Value"
        headers[3].text = "Status"
        headers[4].text = "Remarks"

        for result in reconciliation_results:

            row = table.add_row().cells

            row[0].text = str(result.field_name)
            row[1].text = str(result.database_value)
            row[2].text = str(result.extracted_value)
            row[3].text = str(result.status)
            row[4].text = str(result.remarks)

    # ---------------------------------------------------------
    # Public API
    # ---------------------------------------------------------

    def write(
        self,
        company_name: str,
        fiscal_year: int,
        sections: list[CAMSection],
        output_path: str,
        reconciliation_results: list[ReconciliationResult],
    ) -> str:
        """
        Generate the final CAM document.
        """

        self._add_title(
            company_name,
            fiscal_year,
        )

        for section in sections:

            self._add_section(section)

            # Render Data Consistency Review as a table
            if (
                section.title.lower()
                == "data consistency review"
            ):
                self._add_reconciliation_table(
                    reconciliation_results
                )

        self.document.save(output_path)

        return output_path